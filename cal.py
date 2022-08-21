#! /usr/bin/env python3
"""
Author: neagle2009
Purpose: 生成计算训练算式
!!! 文件名不能与包名相同, 否则会把文件当成包引入, 出现各种属性找不到 !!!
"""

from ctypes.wintypes import RGB
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING

import argparse
import random
import datetime

level = 1
rangeType = 2


def getArgs():
    """ 参数处理 """

    parser = argparse.ArgumentParser(description='加减法算式生成')
    parser.add_argument('-l',
                        '--level',
                        metavar='Level',
                        type=int,
                        default=1,
                        help='难度控制 1: 常规(退位进位都有); 2: 困难模式, 仅退位,进位')
    parser.add_argument('-c',
                        '--count',
                        metavar='Count',
                        type=int,
                        help='生成算式数目')
    parser.add_argument('-t',
                        '--type',
                        metavar='Type',
                        type=int,
                        default=3,
                        help='生成类型, 1: 加法; 2: 减法; 3: 加减混合;')
    parser.add_argument('-r',
                        '--range',
                        help='生成算式范围1: 一位数(TODO); 2: 两位数; 3: 三位数;',
                        metavar='int',
                        type=int,
                        default=2)
    return parser.parse_args()


def GetPlusCal(num):
    """获取加法算式"""

    list = []
    if rangeType == 3:
        rangeMin = 100
        rangeMax = 999
    elif rangeType == 2:
        rangeMin = 10
        rangeMax = 99
    else:
        rangeMin = 10
        rangeMax = 99

    while len(list) < num:
        a = random.randint(rangeMin, rangeMax)
        b = random.randint(rangeMin, rangeMax)
        if (a + b) > (rangeMax + 1):
            continue

        # 进位控制
        p1 = a % 10
        p2 = b % 10
        if (level == 2) and ((p1 + p2) <= 10):
            continue

        result = str(a) + ' + ' + str(b) + ' = '
        if result in list:
            continue
        list.append(result)
    return list


def GetMinusCal(num):
    """获取减法算式"""

    list = []
    if rangeType == 3:
        rangeMin = 100
        rangeMax = 999
    elif rangeType == 2:
        rangeMin = 10
        rangeMax = 99
    else:
        rangeMin = 10
        rangeMax = 99
    while len(list) < num:
        a = random.randint(rangeMin, rangeMax)
        b = random.randint(a, rangeMax)
        if (b - a) <= 10:
            continue

        # 退位控制
        p1 = a % 10
        p2 = b % 10
        if (level == 2) and (p2 >= p1):
            continue

        result = str(b) + ' - ' + str(a) + ' = '
        if result in list:
            continue

        list.append(result)
    return list


def GetMixedCal(num):
    """获取混合算式算式"""

    plus = num // 2
    minus = num - plus
    listP = GetPlusCal(plus)
    listM = GetMinusCal(minus)
    resultList = listP + listM
    resultList.sort()
    return resultList


def printResult(list):
    """ 结果打印(TODO: 直接生成打印格式) """

    i = 0
    print("Count: ", len(list))
    for i in range(0, len(list)):
        print(str(i + 1) + '.', list[i])


def saveDocx(args, list):
    """ 结果保存到docx """

    document = Document()
    document.styles['Normal'].font.name = u'Dejavu Sans Mono for Powerline'
    #字号对照 link: https://www.jianshu.com/p/8f15e3f2f9e6
    document.styles['Normal'].font.size = Pt(14)
    document.styles['Normal']._element.rPr.rFonts.set(
        qn('w:eastAsia'), u'Dejavu Sans Mono for Powerline')

    title = ''
    if args.type == 1:
        title = '加法'
    if args.type == 2:
        title = '减法'
    if args.type == 3:
        title = '加减混合'

    header = document.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.style.font.name = u'Dejavu Sans Mono for Powerline'
    header.style.font.size = Pt(12)
    header.style.font.color.rgb = RGBColor(165, 165, 165)
    header.text = '____年__月__日' + ' 得分: ___' + ' 用时: _____' + ' 评语: _________'
    footer = document.sections[0].footer.paragraphs[0]
    footer.style.font.name = u'Dejavu Sans Mono for Powerline'
    footer.style.font.size = Pt(12)
    footer.style.font.color.rgb = RGBColor(165, 165, 165)   # 灰色  !打印机打印时选择灰度打印, 否则偏蓝
    footer.text = title + '算式(共' + str( len(list)) + '题)    ' + datetime.datetime.now().strftime( '%Y-%m-%d %H:%M:%S')

    #分栏
    document.sections[0]._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')

    for i in range(0, len(list)):
        paragraph = document.add_paragraph(list[i], style='List Number')
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = None  # 段前距
        paragraph_format.space_after = None  # 段后距
        paragraph_format.line_spacing_rule = WD_ALIGN_PARAGRAPH.JUSTIFY  # 设置行距规则
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    x = datetime.datetime.now()
    filename = title + x.strftime("%Y%m%d%H%M%S") + '.docx'
    print('Save to: ' + filename)
    document.save(filename)


def main():
    ''' main function '''

    args = getArgs()
    global level, rangeType
    level = args.level
    rangeType = args.range
    if args.type == 1:
        list = GetPlusCal(args.count)
    elif args.type == 2:
        list = GetMinusCal(args.count)
    else:
        list = GetMixedCal(args.count)
    random.shuffle(list)
    printResult(list)
    saveDocx(args, list)

    x = datetime.datetime.now()
    print(x.strftime("%Y-%m-%d %H:%M:%S"))


if __name__ == '__main__':
    main()
